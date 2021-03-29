#!/usr/bin/python3
# -*- coding: utf-8 -*-

if 0:
    import numpy as np

    def npSum():
        a = np.array([0,1,2,3,4])
        b = np.array([9,8,7,6,5])
        c = a ** 2 + b ** 3
        return c

    print(npSum())

if 0:
    from PyPDF2 import PdfFileReader, PdfFileMerger

    merger = PdfFileMerger()
    input1 = open("Power_Management_Save_Restore.pdf", "rb")
    input2 = open("RPC_programming_interface_v2.0.pdf", "rb")
    # add page 0 ~ 2, total 3
    merger.append(fileobj=input1, pages=(0,3))
    # add page 0, position 2, total 1
    merger.merge(position=2, fileobj=input2, pages=(0,1))
    output = open("output.pdf", "wb")
    merger.write(output)

if 0:
    import docx

    doc = docx.Document()
    doc.add_heading("Document Title", 0)
    p = doc.add_paragraph("A plain paragraph having some ")
    doc.add_page_break()
    doc.save("demo.docx")

if 0:
    import requests

    r = requests.get("https://api.github.com/user", auth=("user", "pass"))
    print(r.status_code)
    print(r.text)

if 0:
    import wx

    app = wx.App(False)
    frame = wx.Frame(None, wx.ID_ANY, "Hello World")
    frame.Show(True)
    app.MainLoop()

if 0:
    import tabulate

    data = [["北京理工大学", "985", 2000],
            ["清华大学", "985", 3000],
            ["大连理工大学", "985", 4000],
            ["深圳大学", "211", 2000],
            ["沈阳大学", "省本", 2000]]
    print(tabulate.tabulate(data, tablefmt="grid"))
